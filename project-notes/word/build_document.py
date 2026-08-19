from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PROJECT_NOTES = Path(__file__).resolve().parents[1]
IMG = PROJECT_NOTES / "assets" / "images"
# Use a distinct output while the earlier document may still be open in Word.
OUT = Path(__file__).resolve().parent / "DCKA_網站與AI助教技術架構_安全強化版.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "ECFDF5"
GREEN = "166534"
LIGHT_RED = "FEF2F2"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "111827"


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "PingFang TC")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 頁")
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def set_running_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("DCKA 技術分享｜網站部署與 AI 助教")
    set_run_font(r, size=9, color=MUTED, bold=True)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_para(doc, text="", *, bold=False, italic=False, size=11, color=BLACK, align=None, before=0, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + "　")
    set_run_font(r, bold=True, color=color)
    r = p.add_run(text)
    set_run_font(r, color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc, filename, caption, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(IMG / filename), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=9, color=MUTED, italic=True)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    r = p.add_run(text)
    set_run_font(r, name="Menlo", size=8.5, color=BLACK)


def add_components_table(doc):
    rows = [
        ("內容層", "Markdown / MkDocs", "保存教材並產生靜態網站與 content.json", "本機 → Pages"),
        ("互動層", "chatbot.js", "保存 Browser Session、傳送問題與 History、渲染回答", "瀏覽器"),
        ("Delivery", "GitHub Actions", "分別建置 GitHub Pages 與部署 Cloud Run", "GitHub"),
        ("Build", "Cloud Build / Registry", "以 dcka-cloud-build（run.builder）建置並保存 container image", "Google Cloud"),
        ("API 層", "FastAPI / google-genai", "CORS、輸入與速率限制、保管 Key、呼叫模型", "Cloud Run"),
        ("模型層", "Gemini 3.5 Flash", "依 System Instruction 與 contents 產生回答", "Google API"),
        ("資料層", "Cloud Firestore", "遮罩匿名問答，保存 created_at 與 90 天 expires_at", "Google Cloud"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [1300, 2000, 4260, 1800]
    set_table_geometry(table, widths)
    headers = ("層級", "技術", "主要責任", "執行位置")
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=NAVY, bold=True)
    mark_header_row(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for i, text in enumerate(values):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 2 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_run_font(r, size=9)
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_validation_table(doc):
    rows = [
        ("GitHub Pages JavaScript", "新版 Cloud Run URL 已上線", "通過"),
        ("Cloud Run GET /", "HTTP 200", "通過"),
        ("POST /api/chat", "HTTP 200；部署驗證約 1.73 秒", "通過"),
        ("CORS", "只允許 GitHub Pages 與 localhost；本機測試通過", "本機通過"),
        ("API 防護", "1 MiB body、4000 字訊息、20 則 History、20 req/min", "本機通過"),
        ("錯誤隔離", "Client 只收到一般錯誤；完整例外留在 Logging", "本機通過"),
        ("Chatbot 實際問答", "繁中回答、文章連結與程式碼顯示正常", "通過"),
        ("Firestore chat_logs", "遮罩、expires_at 與 background write 自動測試通過", "本機通過"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [2600, 4960, 1800]
    set_table_geometry(table, widths)
    for i, text in enumerate(("檢查項目", "結果", "狀態")):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=NAVY, bold=True)
    mark_header_row(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for i, text in enumerate(values):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 2 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, size=9, color=GREEN if i == 2 else BLACK, bold=(i == 2))
        for cell in cells:
            set_cell_margins(cell)


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    configure_styles(doc)
    set_running_header_footer(doc)
    doc.core_properties.title = "DCKA 網站與 AI 助教技術架構"
    doc.core_properties.subject = "MkDocs、GitHub Pages、Cloud Run、FastAPI 與 Gemini 系統介接紀錄"
    doc.core_properties.author = "Charles Cao"

    # Editorial cover (named cover override for the compact reference guide).
    add_para(doc, "TECHNICAL FIELD GUIDE", bold=True, size=10, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, before=70, after=18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("DCKA 網站與 AI 助教")
    set_run_font(r, size=30, color=NAVY, bold=True)
    add_para(doc, "技術架構與系統介接紀錄", size=18, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=26)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    mascot = p.add_run().add_picture(str(ROOT / "docs" / "assets" / "images" / "chatbot-mascot.png"), width=Inches(1.25))
    mascot._inline.docPr.set("descr", "DCKA 學習筆記 AI 助教機器人吉祥物")
    mascot._inline.docPr.set("title", "AI 助教機器人吉祥物")
    add_para(doc, "Markdown → MkDocs → GitHub Pages", bold=True, size=12, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "Browser → Cloud Run / FastAPI → Gemini + Firestore", bold=True, size=12, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=48)
    add_para(doc, "整理：Charles Cao", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "版本：2026-08-19（Asia/Taipei）", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    page_break(doc)
    doc.add_heading("閱讀摘要", level=1)
    add_callout(doc, "一句話說明", "這是一個以 MkDocs 與 GitHub Pages 發布的靜態課程網站；AI 助教由 Cloud Run 呼叫 Gemini，並將匿名問答紀錄持久化到 Firestore。", fill=LIGHT_BLUE)
    add_para(doc, "本文件用於個人技術紀錄與同事分享，重點是理解資料如何從 Markdown 變成網站，以及一次 Chatbot 問答如何跨越瀏覽器、Cloud Run、Gemini 與 Firestore。")
    doc.add_heading("閱讀導覽", level=2)
    for text in (
        "專案整體架構與元件分工",
        "GitHub Pages 前端部署流程",
        "AI 助教的文件上下文與 API 介接",
        "實際網站及 Chatbot 操作畫面",
        "線上驗證結果、安全性與維護建議",
    ):
        add_number(doc, text)
    doc.add_heading("目前正式狀態", level=2)
    add_bullet(doc, "前端：GitHub Pages，正式網址為 https://caocharles.github.io/dcka-class-notes/")
    add_bullet(doc, "後端：Google Cloud Run，FastAPI 提供 / 與 /api/chat。")
    add_bullet(doc, "模型：Gemini 3.5 Flash，thinking level 設為 low。")
    add_bullet(doc, "資料：Cloud Firestore 的 chat_logs 保存遮罩後匿名問答，預設 90 天保留。")
    add_bullet(doc, "安全：exact-origin CORS、輸入上限、instance-local rate limiting 與一般化錯誤已完成。")
    add_bullet(doc, "部署：前端與後端均有獨立 GitHub Actions workflow；本輪變更 push 後才會進正式環境。")

    page_break(doc)
    doc.add_heading("1. 專案整體架構", level=1)
    add_figure(doc, "overall-architecture.png", "圖 1　System Component & Deployment Architecture", width=6.35)
    add_para(doc, "教材閱讀與 AI 問答採前後端分離。圖中實線表示 Runtime 資料流，虛線表示 Build／Deploy 控制流；GitHub、Google Cloud、Browser 與外部 Gemini API 則是四個主要平台或信任邊界。")
    add_callout(doc, "Current State", "Cloud Run 是 stateless public API；畫面對話與匿名 session_id 位於 Browser sessionStorage，問答稽核紀錄則持久化至 Firestore。System Prompt 已由 Backend 控制；每個 Cloud Run instance 按需快取 content.json 一小時，但仍未使用向量資料庫。", fill=LIGHT_BLUE)
    add_components_table(doc)

    page_break(doc)
    doc.add_heading("2. 專案目錄與責任分工", level=1)
    add_code(doc, "dcka-class-notes/\n├── docs/                  # MkDocs 教材、圖片與附件\n├── backend/               # FastAPI、Dockerfile、uv.lock、測試\n├── hooks/                 # 建置後產生 content.json\n├── overrides/             # MkDocs Material 版型覆寫\n├── .github/workflows/     # Pages 與 Cloud Run 自動部署\n├── mkdocs.yml             # 網站與導覽設定\n└── project-notes/         # 本份技術分享紀錄")
    doc.add_heading("關鍵檔案", level=2)
    add_bullet(doc, "mkdocs.yml：決定 nav、theme、plugins、JavaScript、CSS 與 Hook。")
    add_bullet(doc, "hooks/generate_content.py：掃描 docs/**/*.md 並輸出 site/content.json。")
    add_bullet(doc, "docs/assets/js/chatbot.js：Chatbot UI、History 與精簡 API Request；不保存 System Prompt。")
    add_bullet(doc, "backend/chat_server.py：受控 Prompt、教材快取、Gemini client、Firestore logging 與 /api/chat。")
    add_bullet(doc, ".github/workflows/deploy-pages.yml：前端 MkDocs 建置與 GitHub Pages 部署。")
    add_bullet(doc, ".github/workflows/deploy-backend.yml：後端 GitHub Actions 部署。")
    add_callout(doc, "範圍提醒", "project-notes/ 位於 docs/ 外，不會出現在 MkDocs 網站，也不會進入 content.json。", fill=LIGHT_GREEN, color=GREEN)

    page_break(doc)
    doc.add_heading("3. GitHub Pages 網頁部署", level=1)
    add_figure(doc, "github-pages-deployment.png", "圖 2　Frontend & Backend Delivery Architecture", width=6.35)
    doc.add_heading("目前實際方式", level=2)
    add_code(doc, "uv run mkdocs serve\ngit push origin main")
    add_para(doc, "第一個指令在本機預覽；網站相關變更 push 到 main 後，deploy-pages.yml 會依 uv.lock 建置 site/、產生 content.json、上傳 Pages Artifact 並發布。mkdocs gh-deploy 僅保留為緊急 fallback。")
    add_callout(doc, "Repository 設定", "Pages → Build and deployment → Source 已選擇 GitHub Actions；正式發布使用 Pages Artifact。", fill=LIGHT_GREEN, color=GREEN)
    add_callout(doc, "部署差異", "前端 artifact 是 GitHub Pages 的 HTML／CSS／JS／content.json bundle；後端 artifact 是 Cloud Build 建置並保存於 Artifact Registry 的 container image。", fill=LIGHT_BLUE)

    page_break(doc)
    doc.add_heading("4. 網站完成畫面", level=1)
    add_figure(doc, "website-home.png", "圖 3　DCKA 學習筆記公開首頁（2026-08-19）", width=6.35)
    add_para(doc, "使用者可以從頂部導覽切換 Docker／Kubernetes 章節，使用全文搜尋、深色模式與 GitHub 連結；右下角機器人圖示則是 AI 助教入口。")
    doc.add_heading("新增一篇文章", level=2)
    for text in (
        "在 docs/ 新增 Markdown 與需要的圖片。",
        "在 mkdocs.yml 的 nav 加入文章路徑。",
        "執行 mkdocs serve 檢查排版與連結。",
        "commit 並 push main，保存原始碼版本。",
        "push main，讓 deploy-pages.yml 自動建置與發布。",
        "開啟正式網址確認文章與 content.json。",
    ):
        add_number(doc, text)

    page_break(doc)
    doc.add_heading("5. AI 助教系統介接", level=1)
    add_figure(doc, "ai-chatbot-integration.png", "圖 4　AI Assistant Runtime Interaction Architecture", width=6.35)
    add_callout(doc, "核心安全設計", "GEMINI_API_KEY 與 Firestore IAM 只存在 Server side；System Prompt 原始碼仍公開，但執行控制權在 Backend，Browser 無法覆寫，也不會直接存取 Firestore。", fill=LIGHT_GREEN, color=GREEN)
    add_para(doc, "GitHub Pages 與 Cloud Run 是不同 Origin，因此 JSON POST 前會進行 CORS preflight。Backend 只接受正式 Pages 與 localhost Origin，並限制 Body、訊息、History 與速率；同步 worker 呼叫 Gemini 後先回傳一般化 Response，再由 BackgroundTasks 遮罩並寫入 Firestore chat_logs。")

    page_break(doc)
    doc.add_heading("6. Chatbot 使用方式", level=1)
    add_para(doc, "點擊右下角機器人圖示，即可開啟「學習筆記小幫手」。前端不再下載教材；使用者送出問題後，由 Backend 按需讀取並快取 content.json。")
    add_figure(doc, "chatbot-open-focus.png", "圖 5　AI 助教開啟與歡迎畫面", width=3.7)
    add_para(doc, "使用者送出問題後，畫面會顯示思考中狀態；回應完成後，Markdown 會轉成 HTML，並保留文章連結及程式碼複製功能。")
    add_figure(doc, "chatbot-success-focus.png", "圖 6　實際成功問答與課程文章連結", width=3.7)

    page_break(doc)
    doc.add_heading("7. 一次問答的完整流程", level=1)
    steps = (
        "Browser 從 GitHub Pages 取得網站 HTML、CSS 與 JavaScript assets。",
        "對話歷史與匿名 session_id 放在 sessionStorage；Browser 不保存教材或 System Prompt。",
        "使用者輸入問題，Browser 只組合 session_id、history 與 message。",
        "跨 Origin 呼叫先完成 OPTIONS CORS preflight，再 POST /api/chat。",
        "FastAPI 驗證 1 MiB Body、拒絕額外欄位並套用 rate limit；教材快取不存在或超過一小時時，按需讀取 content.json。",
        "Backend 將固定回答規則、信任邊界與完整教材組成 System Instruction，並將 History 轉換成 google.genai.types.Content。",
        "同步 endpoint 由 FastAPI worker thread 執行 google-genai==2.18.1，呼叫 Gemini 3.5 Flash。",
        "FastAPI 先回傳 {text: ...} 或一般化 4xx／5xx；完整例外只寫入 Cloud Logging。",
        "Response 後由 BackgroundTasks 遮罩問答，加入 created_at 與 90 天 expires_at，再寫入 chat_logs。",
        "Firestore 寫入失敗只進 Cloud Logging，不會改變原本的聊天結果。",
        "Browser 使用 fixBrokenLinks() 與 marked.parse() 更新 DOM，再保存 Session。",
    )
    for text in steps:
        add_number(doc, text)
    doc.add_heading("API 請求", level=2)
    add_code(doc, '{\n  "session_id": "7e930f52-...",\n  "history": [{"role": "user", "parts": [{"text": "什麼是 Docker？"}]}],\n  "message": "Kubernetes 的角色是什麼？"\n}')
    doc.add_heading("API 回應", level=2)
    add_code(doc, '{"text": "Kubernetes 是用來部署、調度與管理容器工作負載的編排平台。"}')
    doc.add_heading("Firestore chat_logs", level=2)
    add_code(doc, '{\n  "session_id": "匿名對話 UUID",\n  "question": "遮罩後問題",\n  "answer": "遮罩後回答或 null",\n  "model": "gemini-3.5-flash",\n  "latency_ms": 3700,\n  "status": "success",\n  "error": null,\n  "created_at": "SERVER_TIMESTAMP",\n  "expires_at": "建立時間 + 90 天"\n}')
    add_callout(doc, "RAG 現況", "Backend 會把全站文件直接加入 System Instruction，仍是 full-context RAG；這次改善 Prompt 控制權與 Browser Request 大小，但 Gemini Token 用量不變。", fill=LIGHT_BLUE)

    page_break(doc)
    doc.add_heading("8. GitHub Actions 與 Cloud Run", level=1)
    add_para(doc, "後端 CI/CD 與前端部署是兩條不同流程。當 backend/** 變更並推到 main，deploy-backend.yml 會自動執行：")
    for text in (
        "checkout GitHub repository。",
        "GitHub Actions 取得 OIDC token，透過 Workload Identity Federation 換取短效 GCP 憑證。",
        "以 github-actions-deployer 專用帳號取得 Cloud Run source deploy 權限。",
        "設定 gcloud CLI 與專案。",
        "執行 gcloud run deploy --source backend，指定 dcka-cloud-build 建置並掛載 dcka-chatbot-runtime。",
        "Cloud Build 以 roles/run.builder 依 Dockerfile 建置 image，Artifact Registry 保存 managed artifact。",
        "Cloud Run 建立新 revision 並切換流量。",
        "部署時注入 GEMINI_API_KEY 環境變數。",
    ):
        add_number(doc, text)
    doc.add_heading("GitHub Variables 與 Secret", level=2)
    add_bullet(doc, "Variables：GCP_PROJECT_ID、GCP_WIF_PROVIDER、GCP_DEPLOYER_SA、GCP_BUILD_SA、GCP_RUNTIME_SA。")
    add_bullet(doc, "Secret：GEMINI_API_KEY，用於呼叫 Gemini API。")
    add_callout(doc, "無長效部署金鑰", "WIF Provider 只信任 CaoCharles/dcka-class-notes 的 main branch；repository 不保存 GCP Service Account JSON。", fill=LIGHT_GREEN, color=GREEN)
    add_bullet(doc, "dcka-cloud-build 只具備 roles/run.builder；deployer 透過 roles/iam.serviceAccountUser 啟動這個建置身分。")
    add_callout(doc, "不可提交", "GEMINI_API_KEY 不得寫進 repository、Markdown、Draw.io 或 Word 文件。", fill=LIGHT_RED, color=RED)
    doc.add_heading("Firestore runtime dependency", level=2)
    add_bullet(doc, "Firestore 使用 Native mode，資料位置與 Cloud Run 同為 asia-east1。")
    add_bullet(doc, "Cloud Run 使用 dcka-chatbot-runtime，且只以 roles/datastore.user 存取 Firestore。")
    add_bullet(doc, "Firestore 是一次性建立的 persistent dependency，不會隨 Cloud Run revision 重建。")
    add_bullet(doc, "chat_logs.expires_at TTL policy 已啟用並為 ACTIVE；若開放他人查閱，管理者以專用群組取得 roles/datastore.viewer。")

    page_break(doc)
    doc.add_heading("9. 驗證結果", level=1)
    add_para(doc, "驗證日期：2026-08-19（Asia/Taipei）。安全與 Backend-owned Prompt 已通過後端 13 項自動測試及 MkDocs strict build；既有 GitHub Pages Action、WIF／Cloud Run 部署、正式 CORS、線上問答、Firestore 落盤與 TTL ACTIVE 驗證亦正常。")
    add_validation_table(doc)
    add_para(doc, "簡短 API 測試回應：", bold=True, before=8, after=4)
    add_code(doc, "同學你好！Docker 的用途是將應用程式及其所需的執行環境打包成輕量級的容器，確保程式在任何系統上都能一致、快速地部署與執行。")
    add_callout(doc, "結論", "Gemini、Cloud Run 與 Firestore 的端對端流程已在線上生效；匿名 session_id、問答、模型、延遲與狀態均可正確寫入。", fill=LIGHT_GREEN, color=GREEN)

    page_break(doc)
    doc.add_heading("10. 維護與改善建議", level=1)
    doc.add_heading("本輪已完成", level=2)
    add_bullet(doc, "GitHub Pages／localhost CORS allowlist、API rate limiting 與輸入大小限制。")
    add_bullet(doc, "一般化 Client Error、Cloud Logging 詳細例外、Firestore 遮罩與 90 天 expires_at。")
    add_bullet(doc, "同步 FastAPI endpoint、response 後 BackgroundTasks Firestore logging。")
    add_bullet(doc, "google-genai==2.18.1、backend/uv.lock 與可重現 Docker build。")
    add_bullet(doc, "GitHub Pages frontend Actions、RAG README 與 Cloud Run Dockerfile 註解。")
    add_bullet(doc, "GitHub Actions Workload Identity Federation 與專用 deployer/build/runtime service accounts；線上部署已驗證成功。")
    add_bullet(doc, "Firestore chat_logs.expires_at TTL policy 已啟用並為 ACTIVE。")
    add_bullet(doc, "System Prompt 已移至 Backend；公開 API 拒絕 system_instruction，教材採一小時 instance-local cache 與 stale fallback。")
    doc.add_heading("中期改善", level=2)
    add_bullet(doc, "文件量增加時改成向量檢索式 RAG，只把相關文章片段送給模型。")
    add_bullet(doc, "跨 Cloud Run instances 的全域 rate limit 改用 Cloud Armor、API Gateway 或集中式計數儲存。")
    add_bullet(doc, "將 GEMINI_API_KEY 改由 Secret Manager 注入，強化 rotation 與稽核。")
    add_bullet(doc, "逐步移除預設 Compute Service Account 的 Project Editor 權限；目前 workflow 已不再使用該帳號。")

    page_break(doc)
    doc.add_heading("11. 同事分享建議腳本", level=1)
    add_para(doc, "建議用 10–15 分鐘完成展示：")
    for text in (
        "先開啟 GitHub Pages，說明教材都是 Markdown。",
        "用整體架構圖建立 GitHub Pages 與 Cloud Run 的分工概念。",
        "展示 docs/、mkdocs.yml、Pages build job 與 Pages Artifact 的前端部署關係。",
        "打開 AI 助教並詢問一個 Docker／Kubernetes 問題。",
        "用技術泳道說明 content.json、匿名 session_id、FastAPI、Gemini 與 Firestore。",
        "強調 API Key 與 Firestore IAM 不會出現在瀏覽器，但公開 API 仍需流量保護。",
        "最後分享目前限制與未來改成向量檢索式 RAG 的方向。",
    ):
        add_number(doc, text)
    add_callout(doc, "分享重點", "這個專案的價值不只是一個網站，而是把內容管理、Git 版本控制、雲端部署與生成式 AI 串成一條可維護的學習流程。", fill=LIGHT_BLUE)
    doc.add_heading("相關位置", level=2)
    add_bullet(doc, "網站：https://caocharles.github.io/dcka-class-notes/")
    add_bullet(doc, "原始碼：https://github.com/CaoCharles/dcka-class-notes")
    add_bullet(doc, "可編輯架構圖：project-notes/assets/diagrams/dcka-system-architecture.drawio")
    add_bullet(doc, "Markdown 技術筆記：project-notes/README.md")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
